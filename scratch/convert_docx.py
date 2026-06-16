import docx
import os
import sys

def docx_to_txt(docx_path, txt_path):
    print(f"Converting {docx_path} to {txt_path}...")
    doc = docx.Document(docx_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            # De-duplicate adjacent identical cells due to merged cells
            cleaned_row = []
            for cell in row_text:
                if not cleaned_row or cleaned_row[-1] != cell:
                    cleaned_row.append(cell)
            fullText.append(" | ".join(cleaned_row))
            
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(fullText))

def main():
    src_dir = "/home/abhishek-vishwakarma/Documents/Resume_Run/extracted_data/[PUB] India_runs_data_and_ai_challenge/India_runs_data_and_ai_challenge"
    dest_dir = "/home/abhishek-vishwakarma/Documents/Resume_Run/scratch"
    os.makedirs(dest_dir, exist_ok=True)
    
    files = ["README.docx", "job_description.docx", "redrob_signals_doc.docx", "submission_spec.docx"]
    for file in files:
        src = os.path.join(src_dir, file)
        dest = os.path.join(dest_dir, file.replace(".docx", ".txt"))
        if os.path.exists(src):
            docx_to_txt(src, dest)
        else:
            print(f"File not found: {src}")

if __name__ == "__main__":
    main()
